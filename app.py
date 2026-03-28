from pydoc import doc
from flask import Flask, flash, render_template, request, redirect, send_file, session
from flask_login import LoginManager, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document as DocxDocument
from flask_mail import Mail, Message
import random
from datetime import datetime, timedelta 
import os

from config import Config
from database import Notification, db, User, Template, Document as DocModel


app = Flask(__name__, static_folder='static')
app.config.from_object(Config)
mail=Mail(app)
app.config["MAIL_DEFAULT_SENDER"] = "satwikdhavale1208@gmail.com"

db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def replace_text(doc, key, value):
    for para in doc.paragraphs:
        for run in para.runs:
            if key in run.text:
                run.text = run.text.replace(key, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))

def send_otp(email, otp):

    msg = Message(
        subject="Your OTP Login Code",
        sender="satwikdhavale1208@gmail.com",  
        recipients=[email],
        body=f"Your OTP is: {otp}"
    )

    mail.send(msg)

# ================= INITIAL SETUP ================= #

with app.app_context():
    db.create_all()

    # Default Users
    if not User.query.first():
        users = [
            User(username="admin", password=generate_password_hash("admin123"), role="admin", email="satwikdhavale1208@gmail.com"),
            User(username="faculty", password=generate_password_hash("123"), role="faculty", email="pritibagal25@gmail.com"),
            User(username="hod", password=generate_password_hash("123"), role="hod", email="hod@example.com"),
            User(username="dean", password=generate_password_hash("123"), role="dean", email="dean@example.com"),
            User(username="student", password=generate_password_hash("123"), role="student", email="dhavalesatwik12@gmail.com"),
        ]
        db.session.add_all(users)
        db.session.commit()

    # Default Templates

with app.app_context():
    db.create_all()

    if not Template.query.first():
        templates = [

        # ---------------- STUDENT ----------------
        Template(
            name="Bonafide Certificate",
            file_path="doc_templates/bonafide_template.docx",
            allowed_roles="student,admin",
            approval_flow="hod,dean"
        ),

        Template(
            name="NOC Request",
            file_path="doc_templates/noc_template.docx",
            allowed_roles="student,admin",
            approval_flow="hod,dean"
        ),

        Template(
            name="Internship Permission Request",
            file_path="doc_templates/internship_request.docx",
            allowed_roles="student,admin",
            approval_flow="hod"
        ),

        Template(
            name="Leave Application",
            file_path="doc_templates/leave_application.docx",
            allowed_roles="student,admin",
            approval_flow="faculty"
        ),

        # ---------------- FACULTY ----------------
        Template(
            name="Recommendation Letter",
            file_path="doc_templates/recommendation.docx",
            allowed_roles="faculty,hod,admin",
            approval_flow="hod"
        ),

        Template(
            name="Internship Approval",
            file_path="doc_templates/internship_approval.docx",
            allowed_roles="faculty,hod,admin",
            approval_flow="hod"
        ),

        Template(
            name="Meeting Minutes",
            file_path="doc_templates/meeting_minutes.docx",
            allowed_roles="faculty,hod,admin",
            approval_flow="hod"
        ),

        Template(
            name="Event Certificate",
            file_path="doc_templates/event_certificate.docx",
            allowed_roles="faculty,hod,admin",
            approval_flow="hod"
        ),

        Template(
            name="Notice Draft",
            file_path="doc_templates/notice_draft.docx",
            allowed_roles="faculty,hod,admin",
            approval_flow="hod"
        ),

        # ---------------- HOD ----------------
        Template(
            name="Official Notice",
            file_path="doc_templates/official_notice.docx",
            allowed_roles="hod,admin",
            approval_flow="dean"
        ),

        Template(
            name="Circular",
            file_path="doc_templates/circular.docx",
            allowed_roles="hod,dean,admin",
            approval_flow="dean"
        ),

        Template(
            name="Department Letter",
            file_path="doc_templates/department_letter.docx",
            allowed_roles="hod,admin",
            approval_flow="dean"
        ),

        Template(
            name="Workload Allocation",
            file_path="doc_templates/workload_allocation.docx",
            allowed_roles="hod,admin",
            approval_flow="dean"
        ),

        # ---------------- DEAN ----------------
        Template(
            name="Policy Letter",
            file_path="doc_templates/policy_letter.docx",
            allowed_roles="dean,admin",
            approval_flow="dean"
        ),

        Template(
            name="Appointment Letter",
            file_path="doc_templates/appointment_letter.docx",
            allowed_roles="dean,admin",
            approval_flow="dean"
        ),

        Template(
            name="Promotion Letter",
            file_path="doc_templates/promotion_letter.docx",
            allowed_roles="dean,admin",
            approval_flow="dean"
        ),
    ]
        
        db.session.add_all(templates)
        db.session.commit()

# ================= AUTH ROUTES ================= #

@app.route("/")
def login():
    return render_template("login.html")


@app.route("/login", methods=["POST"])
def login_post():

    username = request.form["username"]
    password = request.form["password"]

    user = User.query.filter_by(username=username).first()

    if not user:
        return render_template("auth_error.html", message="User not found")

    if not user.active:
        return render_template("auth_error.html", message="Account is deactivated")

    if not check_password_hash(user.password, password):
        return render_template("auth_error.html", message="Incorrect password")

    login_user(user)
    return redirect("/dashboard")

from datetime import datetime
import random

@app.route("/otp_login", methods=["GET", "POST"])
def otp_login():

    if request.method == "POST":

        email = request.form["email"]

        user = User.query.filter_by(email=email).first()

        if not user:
            return render_template("otp_login.html", error="User not found")

        # Generate OTP
        otp = str(random.randint(100000, 999999))

        user.otp = otp
        user.otp_created_at = datetime.utcnow()
        db.session.commit()

        # Save email in session
        session["otp_email"] = user.email

        # Send OTP
        send_email(user.email, "Your OTP", f"Your OTP is {otp}")

        print("OTP:", otp)
        
        return redirect("/enter_otp")

    return render_template("otp_login.html")

@app.route("/verify_otp", methods=["POST"])
def verify_otp():

    email = session.get("otp_email")

    if not email:
        return "Session expired"

    user = User.query.filter_by(email=email).first()

    entered_otp = request.form["otp"]

    if user.otp != entered_otp:
        return render_template("verify_otp.html", error="Invalid OTP")

    if datetime.utcnow() > user.otp_created_at + timedelta(minutes=2):
        return render_template("verify_otp.html", error="OTP expired")

    login_user(user)

    return redirect("/dashboard")


@app.route("/resend_otp")
def resend_otp():

    email = session.get("otp_email")

    if not email:
        return "Session expired. Please login again."

    user = User.query.filter_by(email=email).first()

    if not user:
        return "User not found"

    # Generate new OTP
    otp = str(random.randint(100000, 999999))

    user.otp = otp
    user.otp_created_at = datetime.utcnow()
    db.session.commit()

    send_email(
        user.email,
        "OTP Resent",
        f"Your new OTP is: {otp}"
    )


    flash("OTP resent. Please enter the new OTP.", "success")

    return redirect("/enter_otp")

@app.route("/enter_otp")
def enter_otp():
    return render_template("verify_otp.html")

@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect("/")

import re

@app.route("/register", methods=["GET","POST"])
def register():

    if request.method == "POST":

        username = request.form["username"]
        password = request.form["password"]
        confirm_password = request.form["confirm_password"]
        email = request.form["email"]
        role = request.form["role"]

        # ===== VALIDATIONS =====

        # Email format
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            flash("Invalid email format", "error")
            return redirect("/register")

        # Password match
        if password != confirm_password:
            flash("Passwords do not match", "error")
            return redirect("/register")

        # Password strength
        if len(password) < 6:
            flash("Password must be at least 6 characters", "error")
            return redirect("/register")

        # Duplicate email
        if User.query.filter_by(email=email).first():
            flash("Email already registered", "error")
            return redirect("/register")

        # Duplicate username
        if User.query.filter_by(username=username).first():
            flash("Username already exists", "error")
            return redirect("/register")

        # ===== CREATE USER =====
        hashed_password = generate_password_hash(password)

        user = User(
            username=username,
            password=hashed_password,
            role=role,
            email=email
        )

        db.session.add(user)
        db.session.commit()

        flash("Account created successfully!", "success")
        return redirect("/")

    return render_template("register.html")

# ================= DASHBOARD ================= #

@app.route("/dashboard")
@login_required
def dashboard():

    query = DocModel.query

    title = request.args.get("title")
    department = request.args.get("department")
    date = request.args.get("date")
    role = request.args.get("role")
    user = request.args.get("user")

    if title:
        query = query.filter(DocModel.title.contains(title))

    if department:
        query = query.filter(DocModel.department.contains(department))

    if date:
        query = query.filter(DocModel.created_at.contains(date))

    if role:
        query = query.join(User, DocModel.created_by == User.username)\
                 .filter(User.role == role)

    if user:
        query = query.filter(DocModel.created_by.contains(user))

    documents = query.all()

    templates = Template.query.all()

    notifications = Notification.query.filter_by(user_id=current_user.id).all()

    return render_template(
        "dashboard.html",
        documents=documents,
        templates=templates,
        notifications=notifications
    )


@app.route("/profile")
@login_required
def profile():
    notifications = Notification.query.filter_by(user_id=current_user.id).all()
    return render_template("profile.html", notifications=notifications)


@app.route("/change_password", methods=["GET", "POST"])
@login_required
def change_password():
    notifications = Notification.query.filter_by(user_id=current_user.id).all()

    if request.method == "POST":
        current_password = request.form["current_password"]
        new_password     = request.form["new_password"]
        confirm_password = request.form["confirm_password"]

        if not check_password_hash(current_user.password, current_password):
            flash("Current password is incorrect", "error")
            return redirect("/change_password")

        if new_password != confirm_password:
            flash("New passwords do not match", "error")
            return redirect("/change_password")

        if len(new_password) < 6:
            flash("Password must be at least 6 characters", "error")
            return redirect("/change_password")

        current_user.password = generate_password_hash(new_password)
        db.session.commit()

        flash("Password updated successfully!", "success")
        return redirect("/profile")

    return render_template("change_password.html", notifications=notifications)


@app.route("/admin_dashboard")
@login_required
def admin_dashboard():

    if current_user.role != "admin":
        return "Access Denied", 403

    total_documents = DocModel.query.count()

    approved_documents = DocModel.query.filter_by(status="approved").count()

    pending_documents = DocModel.query.filter(
        DocModel.status != "approved",
        DocModel.status != "rejected"
    ).count()

    rejected_documents = DocModel.query.filter_by(status="rejected").count()

    total_users = User.query.count()

    students = User.query.filter_by(role="student").count()
    faculty = User.query.filter_by(role="faculty").count()
    hod = User.query.filter_by(role="hod").count()
    dean = User.query.filter_by(role="dean").count()

    return render_template(
        "admin_dashboard.html",
        total_documents=total_documents,
        approved_documents=approved_documents,
        pending_documents=pending_documents,
        rejected_documents=rejected_documents,
        total_users=total_users,
        students=students,
        faculty=faculty,
        hod=hod,
        dean=dean
    )


# ================= Calendar View ================= #
@app.route("/calendar")
@login_required
def calendar():
    if current_user.role == "admin":
        docs = DocModel.query.all()
    else:
        docs = DocModel.query.filter_by(created_by=current_user.username).all()

    stats = {}
    for doc in docs:
        if doc.created_at:
            date_key = doc.created_at.strftime("%Y-%m-%d")
            if date_key not in stats:
                stats[date_key] = {"sent": 0, "approved": 0, "pending": 0, "docs": []}

            stats[date_key]["sent"] += 1

            if doc.status == "approved":
                stats[date_key]["approved"] += 1
            elif doc.status != "rejected":
                stats[date_key]["pending"] += 1

            stats[date_key]["docs"].append({
                "title": doc.title,
                "template_name": doc.template_name,
                "created_by": doc.created_by,
                "status": doc.status
            })

    events = [
        {"date": date, **s}
        for date, s in stats.items()
    ]

    notifications = Notification.query.filter_by(user_id=current_user.id).all()

    return render_template("calendar.html", events=events, notifications=notifications)

# ================= TEMPLATE ACCESS ================= #

@app.route("/create/<int:template_id>")
@login_required
def create_document(template_id):

    template = Template.query.get_or_404(template_id)
    roles = template.allowed_roles.split(",")

    if current_user.role not in roles:
        return "Access Denied", 403

    return render_template("create_document.html", template=template)


# ================= SUBMIT REQUEST ================= #

@app.route("/submit_request", methods=["POST"])
@login_required
def submit_request():

    title = request.form["title"]
    department = request.form["department"]
    template_name = request.form["template_name"]

    role = current_user.role

    # Decide first approval level
    if role == "student":
        status = "pending_faculty"

    elif role == "faculty":
        status = "pending_hod"

    elif role == "hod":
        status = "pending_dean"

    else:
        status = "approved"

    new_doc = DocModel(
        title=title,
        template_name=template_name,
        department=department,
        created_by=current_user.username,
        status=status
    )

    db.session.add(new_doc)
    db.session.commit()

    if current_user.role == "student":

        faculty_users = User.query.filter_by(role="faculty").all()

        for user in faculty_users:
            send_email(
                user.email,
                "New Request Pending Approval",
                f"""
                A new request has been submitted.

                Title: {title}
                Department: {department}

                Please login to approve/reject.
         """
            )
        
        hod_users = User.query.filter_by(role="hod").all()

    if current_user.role == "faculty":

        hod_users = User.query.filter_by(role="hod").all()

        for user in hod_users:
            send_email(
                user.email,
                "New Request Pending Approval",
                f"""
                A new request has been submitted.

                Title: {title}
                Department: {department}

                Please login to approve/reject.
         """
            )

        dean_users = User.query.filter_by(role="dean").all()

    if current_user.role == "hod":

        dean_users = User.query.filter_by(role="dean").all()

        for user in dean_users:
            send_email(
                user.email,
                "New Request Pending Approval",
                f"""
                A new request has been submitted.

                Title: {title}
                Department: {department}

                Please login to approve/reject.
         """
            )

    return redirect("/dashboard")


# ================= APPROVAL SYSTEM ================= #

@app.route("/approve/<int:doc_id>", methods=["POST"])
@login_required
def approve_document(doc_id):

    doc = DocModel.query.get_or_404(doc_id)

    role = current_user.role

    if role == "faculty" and doc.status == "pending_faculty":
        doc.status = "approved"

    elif role == "hod" and doc.status == "pending_hod":
        doc.status = "approved"

    elif role == "dean" and doc.status == "pending_dean":
        doc.status = "approved"

    # File generation logic
    if doc.status == "approved":

        template = Template.query.filter_by(name=doc.template_name).first()

        if not template:
            return "Template not found"

        filename = doc.title.replace(" ", "_") + ".docx"
        filepath = os.path.join(Config.GENERATED_FOLDER, filename)

        # Create blank doc if template fails
        try:
            document = DocxDocument(template.file_path)

            replace_text(document, "{{TITLE}}", doc.title)
            replace_text(document, "{{DEPARTMENT}}", doc.department)
            replace_text(document, "{{DATE}}", str(doc.created_at.date()) if doc.created_at else "")
            replace_text(document, "{{MESSAGE}}", doc.title)
            replace_text(document, "{{AUTHORITY}}", doc.created_by)
        except:
            document = DocxDocument()

        document.add_paragraph(f"Document: {doc.title}")
        document.add_paragraph(f"Generated for: {doc.created_by}")

        document.save(filepath)

        doc.filename = filename

        print("FILE CREATED:", filepath)

    creator = User.query.filter_by(username=doc.created_by).first()

    send_email(
        creator.email,
        "Document Approved ",
        f"""
        Your document has been approved.

        Title: {doc.title}
        """
    )

    db.session.commit()

    return redirect("/dashboard")

@app.route("/reject/<int:doc_id>", methods=["POST"])
@login_required
def reject_document(doc_id):

    doc = DocModel.query.get_or_404(doc_id)
    doc.status = "rejected"

    creator = User.query.filter_by(username=doc.created_by).first()
    send_email(
        creator.email,
        "Document Rejected",
        f"""
        Your document has been rejected.

        Title: {doc.title}
        """
    )
    db.session.commit()
    return redirect("/dashboard")


# ================= DOWNLOAD ================= #

@app.route("/generated_docs/<filename>")
@login_required
def download_file(filename):

    filepath = os.path.join(Config.GENERATED_FOLDER, filename)

    print("DOWNLOAD REQUEST:", filepath)

    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)

    return "File not found", 404

import threading
mail = Mail(app)
def send_email_async(msg):
    with app.app_context():
        mail.send(msg)
def send_email(to, subject, body):

    msg = Message(
        subject=subject,
        recipients=[to],
        body=body
    )
    threading.Thread(target=send_email_async, args=(msg,)).start()
    

# ================= ADMIN ================= #

@app.route("/manage_users")
@login_required
def manage_users():

    if current_user.role != "admin":
        return "Access Denied", 403

    users = User.query.all()
    return render_template("manage_users.html", users=users)


@app.route("/update_role/<int:user_id>", methods=["POST"])
@login_required
def update_role(user_id):

    if current_user.role != "admin":
        return "Access Denied", 403

    user = User.query.get_or_404(user_id)
    user.role = request.form["role"]

    db.session.commit()
    return redirect("/manage_users")

@app.route("/preview/<int:doc_id>")
@login_required
def preview(doc_id):
    doc = DocModel.query.get_or_404(doc_id)
    return render_template("preview.html", doc=doc)

@app.route("/reset_password/<int:user_id>")
@login_required
def reset_password(user_id):

    if current_user.role != "admin":
        return "Access Denied"

    user = User.query.get_or_404(user_id)

    user.password = generate_password_hash("123456")

    db.session.commit()

    return redirect("/manage_users")

@app.route("/toggle_user/<int:user_id>")
@login_required
def toggle_user(user_id):

    if current_user.role != "admin":
        return "Access Denied"

    user = User.query.get_or_404(user_id)

    user.active = not user.active

    db.session.commit()

    return redirect("/manage_users")

@app.route("/delete_user/<int:user_id>")
@login_required
def delete_user(user_id):

    if current_user.role != "admin":
        return "Access Denied"

    user = User.query.get_or_404(user_id)

    db.session.delete(user)

    db.session.commit()

    return redirect("/manage_users")

@app.route("/create_user", methods=["GET","POST"])
@login_required
def create_user():

    if current_user.role != "admin":
        return "Access Denied"

    if request.method == "POST":

        username = request.form["username"]
        password = generate_password_hash(request.form["password"])
        role = request.form["role"]

        new_user = User(
            username=username,
            password=password,
            role=role
        )

        db.session.add(new_user)
        db.session.commit()

        return redirect("/manage_users")

    return render_template("create_user.html")

# ================= RUN ================= #

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)